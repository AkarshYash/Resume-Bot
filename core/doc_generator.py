import os
from fpdf import FPDF
from docx import Document as DocxDocument

def generate_docx(text: str, output_path: str):
    """Generates a structured .docx file from plain text."""
    doc = DocxDocument()
    
    # Simple styling
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue
            
        # Check if it looks like a header (e.g. all caps or strong sections)
        if line.isupper() and len(line) < 50:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = 14
        else:
            doc.add_paragraph(line)
            
    doc.save(output_path)
    return output_path

def generate_pdf(text: str, output_path: str):
    """Generates a clean, professional PDF resume from plain text."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_margins(15, 15, 15)
    pdf.set_font("Helvetica", size=10)
    
    # Common unicode replacements for Helvetica
    replacements = {
        '•': '-', '–': '-', '—': '-', '’': "'", '‘': "'",
        '“': '"', '”': '"', '…': '...', '✔': '-', '✓': '-',
        '⭐': '*', '\t': '    '
    }

    # Line spacing and styling
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            pdf.ln(4)
            continue
            
        for k, v in replacements.items():
            line = line.replace(k, v)
            
        # Encode to latin-1 to drop any remaining unsupported unicode that breaks Helvetica
        line = line.encode('latin-1', 'ignore').decode('latin-1')
        
        # Highlight large section headers
        if line.isupper() and len(line) < 50:
            pdf.ln(2)
            pdf.set_font("Helvetica", style="B", size=12)
            pdf.cell(0, 8, line, ln=1)
            pdf.set_font("Helvetica", size=10)
            pdf.ln(1)
        else:
            # Multi-cell for wrapping text
            try:
                pdf.multi_cell(0, 5, line)
            except Exception as e:
                # Fallback if a line still breaks (e.g. unbroken string longer than page width)
                # Split it manually into chunks of 80 characters
                chunks = [line[i:i+80] for i in range(0, len(line), 80)]
                for chunk in chunks:
                    pdf.cell(0, 5, chunk, ln=1)
            
    pdf.output(output_path)
    return output_path
